import warnings
warnings.filterwarnings("ignore", category=UserWarning)

from flask import Flask, render_template, jsonify, request, url_for
import os
import traceback
from werkzeug.utils import secure_filename
from qa_analyse import *
import torch
#import torchvision
import numpy as np
import json
import logging
from docx import Document
from docx2pdf import convert
import pythoncom

try:
    from fpdf import FPDF
except ImportError:
    import os
    os.system('pip install fpdf')
    from fpdf import FPDF
# ...existing code...

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Définir NumpyEncoder avant de l'utiliser
class NumpyEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        if isinstance(obj, np.float32):
            return float(obj)
        if isinstance(obj, torch.Tensor):
            return obj.detach().cpu().numpy().tolist()
        return super(NumpyEncoder, self).default(obj)

# Configurer Flask pour utiliser le NumpyEncoder
app = Flask(__name__)
app.json_encoder = NumpyEncoder  # Ajouter cette ligne après la création de l'app

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'uploads')
app.static_folder = os.path.join(BASE_DIR, 'static')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max-limit

# Assurer que les dossiers nécessaires existent
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.static_folder, exist_ok=True)

SAVED_TEXT_FOLDER = 'G:/OneDrive/Entreprendre/Actions-4'
ALLOWED_EXTENSIONS = {'pdf'}

#app.config['SAVED_TEXT_FOLDER'] = SAVED_TEXT_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/save-path')
def getSavePath():
    return SAVED_TEXT_FOLDER

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400

        file1 = request.files['file']
        
        if file1.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if not allowed_file(file1.filename):
            return jsonify({'error': 'File type not allowed'}), 400
        
        file1_path = file1.filename[:4]
        if not os.path.exists(app.config['UPLOAD_FOLDER']):
            os.makedirs(app.config['UPLOAD_FOLDER'])

        file = os.path.join(app.config['UPLOAD_FOLDER'],file1.filename)
        file1.save(file)

        return jsonify({
            'path': file,
            'file_dir': file1_path
        })

    except Exception as e:
        print(f"Upload error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/job-details')
def job_details():
    return render_template('job_details.html')

# get_answer
@app.route('/get_job_answer', methods=['POST'])
def extract_job_text():
    try:
        file = request.json.get('path')
        RQ = request.json.get('RQ')
        
        if not file or not RQ:
            return jsonify({'error': 'Missing job file path or question'}), 400

        # Extraction rapide du texte
        text1 = extract_text_from_pdf(file)
        print("text1 ",text1)
        if not text1:
            return jsonify({'error': 'Job text extraction failed'}), 500
       
        # formated pour affichage text
        answer = get_answer(RQ, text1 )
        print("answer ",answer)
        # Ensure consistent formatting of the response
        return jsonify({
            'raw_text': text1,
            'formatted_text': answer
        })

    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'error': str(e)}), 500


""" 
def save_job_text_as_pdf(job_text_data, file_path_full):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    
    for line in job_text_data.split('\n'):
        pdf.multi_cell(0, 10, line)
    
    pdf_output_path = file_path_full.replace('.txt', '.pdf')
    pdf.output(pdf_output_path)
    return pdf_output_path
 """
@app.route('/save-job-text', methods=['POST'])
def save_job_text():
    try:
        pythoncom.CoInitialize()  # Initialize COM library
        job_text_data = request.json.get('job_text_data')
        job_number = request.json.get('job_number')
        path= request.json.get('path')
        if not job_text_data or not job_number:
            logger.error(f"Missing job text data or job number: job_text_data={job_text_data}, job_number={job_number}")
            return jsonify({'error': 'Missing job text data or job number'}), 400

        file_name = f"{job_number}_gpt_request"
        file_path = os.path.join(SAVED_TEXT_FOLDER, job_number)
        file_path_txt = os.path.join(file_path, file_name + ".txt")
        file_path_docx = os.path.join(file_path, file_name + ".docx")
        """ Save txt file"""
        with open(file_path_txt, 'w', encoding='utf-8') as file:
            file.write(job_text_data)
        
        """ Save pdf file"""
        doc = format_text_as_word_style(job_text_data, job_number)
        doc.save(file_path_docx)

        pdf_file_path = file_path_docx.replace('.docx', '.pdf')
        convert(file_path_docx, pdf_file_path)

        return jsonify({'message': 'Job text saved successfully', 'file_path': file_path_txt, 'pdf_file_path': pdf_file_path})

    except Exception as e:
        logger.error(f"Error saving job text: {str(e)}")
        return jsonify({'error': str(e)}), 500

    finally:
        pythoncom.CoUninitialize()  # Uninitialize COM library
        

def format_text_as_word_style(job_text, job_number):
    doc = Document()
    doc.add_heading(f"Job Number: {job_number}", level=1)
    
    for line in job_text.split('\n'):
        if line.startswith('- '):
            doc.add_paragraph(line, style='ListBullet')
        else:
            doc.add_paragraph(line)
    
    return doc

    """--------------------------- Analyse_
    
    """
@app.route('/extract_features', methods=['POST'])
def extract_features(text):
    print("\n2. Extracting features...")
    cv_features = extract_features(cv_text)
    job_features = extract_features(job_text)
    print("\n3. Computing similarity...")
    raw_similarity = cosine_similarity(cv_features, job_features)
    adjusted_similarity = calculate_similarity_score(cv_features, job_features)
        
    print(f"\nResults:")
    raw_similarity_s=f"{raw_similarity:.4f}"
    adjusted_similarity_s=f"{adjusted_similarity:.4f}"
    print(f"Raw similarity score: {raw_similarity_s}")
    print(f"Adjusted similarity score: {adjusted_similarity_s}")
    return jsonify({'Raw_similarity_score': raw_similarity_s, 'Adjusted_similarity_score':adjusted_similarity_s})
    
if __name__ == '__main__':
    app.run(debug=False)